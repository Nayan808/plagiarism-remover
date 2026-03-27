import io
import os
import re
import tempfile
from pathlib import Path

from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from groq import Groq
from mangum import Mangum

load_dotenv()

# ── Paraphraser ──────────────────────────────────────────────────

SYSTEM_PROMPT = (
    "You are a professional paraphrasing assistant. "
    "Rewrite text to be unique and plagiarism-free while:\n"
    "- Preserving the EXACT same meaning and technical accuracy\n"
    "- Keeping approximately the SAME word count (±10%)\n"
    "- Using the same sentence structure and tone\n"
    "- NOT adding headers, bullets, explanations, or extra content\n"
    "Return ONLY the rewritten text with no commentary."
)


def _get_client() -> Groq:
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError(
            "GROQ_API_KEY environment variable is not set. "
            "Get a free key at https://console.groq.com"
        )
    return Groq(api_key=api_key)


def paraphrase_batch(paragraphs: list[str], model: str) -> list[str]:
    """Paraphrase paragraphs using numbered markers for reliable parsing."""
    import time
    non_empty = [(i, p) for i, p in enumerate(paragraphs) if p.strip() and len(p.strip()) >= 10]
    if not non_empty:
        return paragraphs

    # Build numbered input: [1] para1\n[2] para2\n...
    numbered_input = "\n".join(f"[{n+1}] {p.strip()}" for n, (_, p) in enumerate(non_empty))
    prompt = (
        "Paraphrase each numbered paragraph below. "
        "Return ONLY the numbered paragraphs in the exact same order, "
        "each starting with [N] on its own line. "
        "Keep the same length and do not merge or split paragraphs.\n\n"
        + numbered_input
    )

    client = _get_client()
    for attempt in range(4):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
                max_tokens=4096,
            )
            raw = response.choices[0].message.content.strip()

            # Parse [N] markers reliably
            result = list(paragraphs)
            for n, (orig_i, orig_p) in enumerate(non_empty):
                pattern = rf'\[{n+1}\]\s*(.*?)(?=\[{n+2}\]|\Z)'
                m = re.search(pattern, raw, re.DOTALL)
                if m:
                    text = re.sub(r'^["\']|["\']$', "", m.group(1)).strip()
                    result[orig_i] = text if text else orig_p
                else:
                    result[orig_i] = orig_p   # fallback to original
            return result
        except Exception as e:
            if "429" in str(e) and attempt < 3:
                time.sleep(2 ** attempt * 3)
                continue
            raise
    return paragraphs


def paraphrase_text(text: str, model: str = "llama-3.1-8b-instant") -> str:
    """Single paragraph fallback (used for table cells)."""
    import time
    text = text.strip()
    if not text or len(text) < 10:
        return text
    client = _get_client()
    for attempt in range(4):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"Paraphrase this:\n\n{text}"},
                ],
                temperature=0.7,
                max_tokens=2048,
            )
            result = response.choices[0].message.content.strip()
            return re.sub(r'^["\']|["\']$', "", result).strip()
        except Exception as e:
            if "429" in str(e) and attempt < 3:
                time.sleep(2 ** attempt * 3)
                continue
            raise
    return text


# ── AI Detection ─────────────────────────────────────────────────

def _extract_sample(data: bytes, ext: str, max_chars: int = 2000) -> str:
    try:
        if ext == "txt":
            return data.decode("utf-8", errors="ignore")[:max_chars]
        elif ext == "docx":
            from docx import Document
            doc = Document(io.BytesIO(data))
            return " ".join(p.text for p in doc.paragraphs if p.text.strip())[:max_chars]
        elif ext == "pdf":
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            return "".join(page.get_text() for page in doc)[:max_chars]
    except Exception:
        pass
    return ""


def detect_ai_content(text: str, model: str) -> int:
    """Returns 0-100 AI content percentage, or -1 on failure."""
    if not text.strip():
        return -1
    try:
        client = _get_client()
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an AI content detector. "
                        "Respond with ONLY a single integer from 0 to 100 representing "
                        "the percentage of the text that is AI-generated. "
                        "0 = fully human-written, 100 = fully AI-generated. "
                        "No explanation. Just the number."
                    ),
                },
                {"role": "user", "content": text},
            ],
            temperature=0,
            max_tokens=5,
        )
        raw = response.choices[0].message.content.strip()
        m = re.search(r"\d+", raw)
        if m:
            return min(100, max(0, int(m.group())))
    except Exception:
        pass
    return -1


# ── Document Handler ─────────────────────────────────────────────

def _get_para_full_text(para) -> str:
    return "".join(run.text for run in para.runs)


def _set_para_text_preserve_format(para, new_text: str):
    """Replace paragraph text while preserving per-run formatting (bold, italic, font, size)."""
    if not para.runs:
        para.text = new_text
        return

    runs = para.runs
    if len(runs) == 1:
        runs[0].text = new_text
        return

    # Distribute new text proportionally across runs by original character count
    orig_lengths = [len(r.text) for r in runs]
    total = sum(orig_lengths)

    if total == 0:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        return

    pos = 0
    new_len = len(new_text)
    for i, (run, orig_len) in enumerate(zip(runs, orig_lengths)):
        if i == len(runs) - 1:
            run.text = new_text[pos:]
        else:
            chars = round(new_len * orig_len / total)
            end = pos + chars
            # Snap to nearest word boundary
            if end < new_len:
                space = new_text.rfind(" ", pos, end + 20)
                if space > pos:
                    end = space + 1
            end = min(end, new_len)
            run.text = new_text[pos:end]
            pos = end


def process_txt(content: bytes, model: str) -> bytes:
    text = content.decode("utf-8", errors="ignore")
    paragraphs = text.split("\n")
    result = paraphrase_batch(paragraphs, model)
    return "\n".join(result).encode("utf-8")


def process_docx(content: bytes, model: str) -> bytes:
    from docx import Document
    doc = Document(io.BytesIO(content))

    # Batch-paraphrase all body paragraphs in one API call
    paras = doc.paragraphs
    texts = [_get_para_full_text(p) for p in paras]
    paraphrased = paraphrase_batch(texts, model)
    for para, new_text in zip(paras, paraphrased):
        if new_text.strip():
            _set_para_text_preserve_format(para, new_text)

    # Table cells: collect all, batch once
    cell_paras = [para for table in doc.tables for row in table.rows
                  for cell in row.cells for para in cell.paragraphs]
    cell_texts = [_get_para_full_text(p) for p in cell_paras]
    cell_paraphrased = paraphrase_batch(cell_texts, model)
    for para, new_text in zip(cell_paras, cell_paraphrased):
        if new_text.strip():
            _set_para_text_preserve_format(para, new_text)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def _docx_to_pdf(docx_path: str, tmpdir: str, pdf_out: str) -> bytes:
    import subprocess
    import shutil
    lo_path = shutil.which("soffice") or shutil.which("libreoffice")
    if lo_path:
        try:
            subprocess.run(
                [lo_path, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                check=True, capture_output=True, timeout=120,
            )
            out_pdf = os.path.join(tmpdir, Path(docx_path).stem + ".pdf")
            if os.path.exists(out_pdf):
                with open(out_pdf, "rb") as f:
                    return f.read()
        except Exception:
            pass
    with open(docx_path, "rb") as f:
        return f.read()


def process_pdf(content: bytes, model: str) -> bytes:
    from pdf2docx import Converter
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_in = os.path.join(tmpdir, "input.pdf")
        docx_converted = os.path.join(tmpdir, "converted.docx")
        docx_processed = os.path.join(tmpdir, "processed.docx")
        pdf_out = os.path.join(tmpdir, "output.pdf")
        with open(pdf_in, "wb") as f:
            f.write(content)
        cv = Converter(pdf_in)
        cv.convert(docx_converted, start=0, end=None)
        cv.close()
        with open(docx_converted, "rb") as f:
            docx_bytes = f.read()
        processed_bytes = process_docx(docx_bytes, model)
        with open(docx_processed, "wb") as f:
            f.write(processed_bytes)
        return _docx_to_pdf(docx_processed, tmpdir, pdf_out)


# ── FastAPI App ──────────────────────────────────────────────────

app = FastAPI(title="Plagiarism Remover API")

_raw_origins = os.environ.get("ALLOWED_ORIGINS", "*")
ALLOWED_ORIGINS = [o.strip() for o in _raw_origins.split(",")] if _raw_origins != "*" else ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
    expose_headers=["X-Output-Filename", "X-AI-Before", "X-AI-After"],
)

EXTENSION_MAP = {".docx": "docx", ".pdf": "pdf", ".txt": "txt"}
CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":  "application/pdf",
    "txt":  "text/plain; charset=utf-8",
}
ALLOWED_MODELS = {
    "llama-3.1-8b-instant",
    "llama-3.3-70b-versatile",
    "llama-3.3-70b-specdec",
    "meta-llama/llama-4-scout-17b-16e-instruct",
    "qwen/qwen3-32b",
}


@app.get("/")
def root():
    return {"status": "Plagiarism Remover API is running"}


@app.get("/health")
def health():
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        return {"groq": "missing", "error": "GROQ_API_KEY not set"}
    return {"groq": "configured", "models": sorted(ALLOWED_MODELS)}


@app.post("/process")
async def process_file(
    file: UploadFile = File(...),
    model: str = Form(default="llama-3.1-8b-instant"),
):
    filename = file.filename or "upload"
    ext = Path(filename).suffix.lower()
    file_type = EXTENSION_MAP.get(ext)

    if not file_type:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Use .docx, .pdf, or .txt",
        )

    if model not in ALLOWED_MODELS:
        model = "llama-3.1-8b-instant"

    content = await file.read()

    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Max 20 MB.")

    # AI detection — before
    ai_before = detect_ai_content(_extract_sample(content, file_type), model)

    try:
        if file_type == "txt":
            result = process_txt(content, model)
            out_ext = ".txt"
        elif file_type == "docx":
            result = process_docx(content, model)
            out_ext = ".docx"
        else:
            result = process_pdf(content, model)
            out_ext = ".pdf" if result[:4] == b"%PDF" else ".docx"
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")

    # AI detection — after
    ai_after = detect_ai_content(_extract_sample(result, out_ext.lstrip(".")), model)

    stem = Path(filename).stem
    out_filename = f"{stem}_paraphrased{out_ext}"
    content_type = CONTENT_TYPES.get(out_ext.lstrip("."), "application/octet-stream")

    return Response(
        content=result,
        media_type=content_type,
        headers={
            "Content-Disposition": f'attachment; filename="{out_filename}"',
            "X-Output-Filename": out_filename,
            "X-AI-Before": str(ai_before),
            "X-AI-After":  str(ai_after),
        },
    )


# Vercel serverless handler
handler = Mangum(app, lifespan="off")
