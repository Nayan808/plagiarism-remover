import io
import os
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import fitz  # PyMuPDF
from pdf2docx import Converter

from paraphraser import paraphrase_text


# ─────────────────────────────────────────────
#  PLAIN TEXT
# ─────────────────────────────────────────────

def process_txt(content: bytes, model: str) -> bytes:
    text = content.decode("utf-8", errors="ignore")
    paragraphs = text.split("\n")
    output_parts = []
    for para in paragraphs:
        if para.strip():
            output_parts.append(paraphrase_text(para, model))
        else:
            output_parts.append(para)
    return "\n".join(output_parts).encode("utf-8")


# ─────────────────────────────────────────────
#  WORD (.docx)  — preserves formatting
# ─────────────────────────────────────────────

def _get_para_full_text(para) -> str:
    return "".join(run.text for run in para.runs)


def _set_para_text_preserve_format(para, new_text: str):
    """
    Replace paragraph text while keeping the formatting of the first run.
    All subsequent runs are cleared so we don't duplicate text.
    """
    if not para.runs:
        para.text = new_text
        return

    # Keep first run's formatting, put all new text there
    first_run = para.runs[0]
    first_run.text = new_text

    # Clear all other runs
    for run in para.runs[1:]:
        run.text = ""


def process_docx(content: bytes, model: str) -> bytes:
    doc = Document(io.BytesIO(content))

    # Process normal paragraphs
    for para in doc.paragraphs:
        original = _get_para_full_text(para).strip()
        if original:
            paraphrased = paraphrase_text(original, model)
            _set_para_text_preserve_format(para, paraphrased)

    # Process text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = _get_para_full_text(para).strip()
                    if original:
                        paraphrased = paraphrase_text(original, model)
                        _set_para_text_preserve_format(para, paraphrased)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


# ─────────────────────────────────────────────
#  PDF  — convert → process → convert back
# ─────────────────────────────────────────────

def process_pdf(content: bytes, model: str) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_in = os.path.join(tmpdir, "input.pdf")
        docx_converted = os.path.join(tmpdir, "converted.docx")
        docx_processed = os.path.join(tmpdir, "processed.docx")
        pdf_out = os.path.join(tmpdir, "output.pdf")

        # Save uploaded PDF
        with open(pdf_in, "wb") as f:
            f.write(content)

        # PDF → DOCX (preserves layout)
        cv = Converter(pdf_in)
        cv.convert(docx_converted, start=0, end=None)
        cv.close()

        # Paraphrase the DOCX
        with open(docx_converted, "rb") as f:
            docx_bytes = f.read()
        processed_bytes = process_docx(docx_bytes, model)
        with open(docx_processed, "wb") as f:
            f.write(processed_bytes)

        # DOCX → PDF using python-docx + reportlab fallback via fitz
        # We'll use LibreOffice if available, otherwise return docx
        result_bytes = _docx_to_pdf(docx_processed, tmpdir, pdf_out)
        return result_bytes


def _docx_to_pdf(docx_path: str, tmpdir: str, pdf_out: str) -> bytes:
    """Try LibreOffice conversion; fall back to returning the DOCX."""
    import subprocess
    import shutil

    # Try LibreOffice (free, installed separately)
    lo_path = shutil.which("soffice") or shutil.which("libreoffice")
    if lo_path:
        try:
            subprocess.run(
                [lo_path, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                check=True,
                capture_output=True,
                timeout=120,
            )
            out_pdf = os.path.join(tmpdir, Path(docx_path).stem + ".pdf")
            if os.path.exists(out_pdf):
                with open(out_pdf, "rb") as f:
                    return f.read()
        except Exception:
            pass

    # Fallback: return processed DOCX (user can open in Word and save as PDF)
    with open(docx_path, "rb") as f:
        return f.read()
